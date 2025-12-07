# -----------------------------
# Environment / Imports
# -----------------------------
import os
# Disable inotify watcher, force polling (helps on Streamlit Cloud)
os.environ["STREAMLIT_WATCHER_TYPE"] = "poll"
os.environ["WATCHDOG_DISABLE_FILE_WATCHING"] = "true"

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import ttest_ind
import openpyxl
from pathlib import Path
import plotly.express as px

st.set_page_config(layout="wide")

# -----------------------------
# Helpers
# -----------------------------
@st.cache_data(show_spinner=False)
def read_excel_table(path: Path, sheet_name: str, table_name: str):
    """
    Read an Excel 'table' by name from a given sheet.
    If the table is missing, fall back to reading the entire sheet.
    """
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=False)
        sheet = wb[sheet_name]

        if table_name in sheet.tables:
            # Case 1: read the named Excel table range
            table = sheet.tables[table_name]
            ref = table.ref
            min_col, min_row, max_col, max_row = openpyxl.utils.range_boundaries(ref)

            data = []
            for row in sheet.iter_rows(
                min_row=min_row, max_row=max_row,
                min_col=min_col, max_col=max_col,
                values_only=True
            ):
                data.append(row)

            df = pd.DataFrame(data[1:], columns=data[0])
        else:
            # Case 2: no table; read whole sheet using the first row as header
            data = sheet.values
            cols = next(data)
            df = pd.DataFrame(data, columns=cols)

        # Drop unnamed columns
        df = df.loc[:, ~df.columns.astype(str).str.startswith('Unnamed')]

        return df

    except Exception as e:
        st.error(f"❌ Error reading from sheet '{sheet_name}': {e}")
        return None


def cohen_d(x, y):
    nx = len(x)
    ny = len(y)
    pooled_std = np.sqrt(((nx - 1) * np.std(x, ddof=1) ** 2 +
                          (ny - 1) * np.std(y, ddof=1) ** 2) /
                         (nx + ny - 2))
    return (np.mean(x) - np.mean(y)) / pooled_std if pooled_std != 0 else 0


# -----------------------------
# Main
# -----------------------------
def main():
    st.title("📊 Water Polo International Analysis Page")

    # -------------------------
    # Load data from repo root
    # -------------------------
    DATA_PATH = Path(__file__).parent / "Winning_Losing_Teams.xlsx"
    # (Keep this line if you want a quick on-screen check)
    # st.write(f"Looking for file at: {DATA_PATH}")

    if not DATA_PATH.exists():
        st.error(f"❌ File not found at {DATA_PATH}. Did you push it to GitHub?")
        st.stop()

    # Read sheets (table fallback supported)
    df_win = read_excel_table(DATA_PATH, "Winning Teams", "Table1")
    df_loss = read_excel_table(DATA_PATH, "Losing Teams", "Table2")

    if df_win is None or df_loss is None:
        st.error("❌ Could not read data from the Excel file.")
        st.stop()

    # -----------------------------------------
    # 1) Global filters (apply to ALL tabs)
    # -----------------------------------------

    with st.sidebar:
        st.markdown("## 🧭 Global Filters")
        st.markdown("Use these filters to refine data across all tabs.")

        # ---------- Filter by Competition ----------
        with st.expander("🔍 Filter by Competition", expanded=True):
            if 'Competition' in df_win.columns and 'Competition' in df_loss.columns:
                competition_options = sorted(set(df_win['Competition'].dropna()) | set(df_loss['Competition'].dropna()))
                selected_competitions = st.multiselect(
                    "Select Competitions",
                    competition_options,
                    default=competition_options
                )
                if selected_competitions:
                    df_win_filtered = df_win[df_win['Competition'].isin(selected_competitions)].copy()
                    df_loss_filtered = df_loss[df_loss['Competition'].isin(selected_competitions)].copy()
                else:
                    df_win_filtered = df_win.copy()
                    df_loss_filtered = df_loss.copy()
            else:
                st.warning("⚠️ 'Competition' column not found.")
                df_win_filtered = df_win.copy()
                df_loss_filtered = df_loss.copy()

        # ---------- Filter by Match Tier ----------
        with st.expander("🎯 Filter by Match Tier", expanded=False):
            if 'Tier' in df_win.columns and 'Tier' in df_loss.columns:
                with st.popover("ℹ️ Tiering System Explained"):
                    st.markdown("""
                    - **Tier 1 Matches:** Top 7 teams play each other  
                    - **Tier 2 Matches:** One Top 7 team vs. a team outside the Top 7, or two lower teams  
                    - **Tier 1 Teams:** AUS, ESP, GRE, HUN, ITA, NED, USA  
                    - **Tier 2 Teams:** All other teams  
                    """)

                tier_options = sorted(set(df_win['Tier'].dropna()) | set(df_loss['Tier'].dropna()))
                selected_tier = st.radio(
                    "Select Tier of Matches",
                    ["All Matches"] + tier_options,
                    index=0
                )

                if selected_tier != "All Matches":
                    df_win_filtered = df_win_filtered[df_win_filtered['Tier'] == selected_tier].copy()
                    df_loss_filtered = df_loss_filtered[df_loss_filtered['Tier'] == selected_tier].copy()
            else:
                st.warning("⚠️ 'Tier' column not found.")

        # ---------- Filter by Timeframe ----------
        with st.expander("🗓️ Filter by Timeframe", expanded=False):
            if 'Date' in df_win.columns and 'Date' in df_loss.columns:
                df_win['Date'] = pd.to_datetime(df_win['Date'], errors='coerce')
                df_loss['Date'] = pd.to_datetime(df_loss['Date'], errors='coerce')

                min_date = pd.concat([df_win['Date'], df_loss['Date']]).min()
                max_date = pd.Timestamp.today()

                start_date = st.date_input("📆 From", value=min_date, min_value=min_date, max_value=max_date)
                end_date = st.date_input("📅 To", value=max_date, min_value=min_date, max_value=max_date)

                df_win_filtered = df_win_filtered[
                    (df_win_filtered['Date'] >= pd.Timestamp(start_date)) &
                    (df_win_filtered['Date'] <= pd.Timestamp(end_date))
                    ].copy()

                df_loss_filtered = df_loss_filtered[
                    (df_loss_filtered['Date'] >= pd.Timestamp(start_date)) &
                    (df_loss_filtered['Date'] <= pd.Timestamp(end_date))
                    ].copy()

                st.caption(
                    f"Currently showing data between **{start_date.strftime('%d %b %Y')}** and **{end_date.strftime('%d %b %Y')}**."
                )
            else:
                st.warning("⚠️ 'Date' column not found.")

    # ✅ Common numeric columns used throughout
    num_cols = df_win_filtered.select_dtypes(include=np.number).columns.intersection(
        df_loss_filtered.select_dtypes(include=np.number).columns
    )

    # -----------------------------------------
    # 2) Tabs
    # -----------------------------------------
    tabs = st.tabs([
        "🏠 Home",
        "📊 Winning and Losing Differences",
        "📈 Performance Index",
        "🆚 Comparison Table",
        "🔍 Team Breakdown",
        "🔮 Predictive Insights",
        "⚖️ Referee Analysis",
        "🤽‍♀️ Opponent Profile"
    ])

    # -------------------------
    # Tab 0: HOME
    # -------------------------
    with tabs[0]:
        try:
            st.subheader("📌 Global Average Comparison (Winning vs Losing Teams)")

            # Global Summary
            avg_win = df_win_filtered[num_cols].mean()
            avg_loss = df_loss_filtered[num_cols].mean()

            global_df = pd.DataFrame({'Winning Teams': avg_win, 'Losing Teams': avg_loss}).sort_index()

            # Plotly-friendly
            global_df_reset = global_df.reset_index().rename(columns={'index': 'Statistic'})
            global_df_melted = global_df_reset.melt(
                id_vars='Statistic', var_name='Result', value_name='Average Value'
            )

            fig = px.bar(
                global_df_melted,
                y='Statistic',
                x='Average Value',
                color='Result',
                barmode='group',
                orientation='h',
                hover_name='Statistic',
                hover_data={'Average Value': ':.2f', 'Result': True},
                color_discrete_map={'Winning Teams': '#66c2a5', 'Losing Teams': '#fc8d62'},
                height=800
            )
            fig.update_layout(xaxis_title="Average Value", yaxis_title="Statistic")
            st.plotly_chart(fig, use_container_width=True)

        except Exception as e:
            st.error(f"❌ Error rendering Home tab: {e}")

        # --- 🧠 Context-Adjusted Performance Summary (Data-Driven) ---
        st.markdown("### 🧠 Context-Adjusted Performance Summary")

        try:
            # Confirm tier context
            tier_context = "All Matches"
            if "selected_tier" in locals():
                tier_context = selected_tier

            # Identify key metrics based on your actual dataset
            key_stats = [
                stat for stat in [
                    "Goals Scored", "Goals Conceded", "Shots", "SoT",
                    "Exclusions Won", "Exclusions Conceded",
                    "Blocks", "Pass to CF", "GK Save"
                ]
                if stat in df_win_filtered.columns
            ]

            # Compute averages per tier
            df_all = df_win_filtered.copy()
            tier_1_df = df_all[df_all["Tier"] == "Tier 1"] if "Tier" in df_all.columns else pd.DataFrame()
            tier_2_df = df_all[df_all["Tier"] == "Tier 2"] if "Tier" in df_all.columns else pd.DataFrame()

            avg_all = df_all[key_stats].mean().round(1)
            avg_t1 = tier_1_df[key_stats].mean().round(1) if not tier_1_df.empty else None
            avg_t2 = tier_2_df[key_stats].mean().round(1) if not tier_2_df.empty else None

            summary_lines = []

            # Helper function to describe Tier 1 vs Tier 2 difference
            def describe_change(stat, t1, t2):
                diff = t1 - t2
                if abs(diff) < 0.3:
                    return None
                direction = "higher" if diff > 0 else "lower"
                return f"- **{stat}** averages are **{direction}** in Tier 1 matches ({t1:.1f}) vs Tier 2 ({t2:.1f})."

            # --- Generate contextual insights dynamically ---
            if avg_t1 is not None and avg_t2 is not None:
                for stat in key_stats:
                    text = describe_change(stat, avg_t1[stat], avg_t2[stat])
                    if text:
                        summary_lines.append(text)

            # Write a short tier-specific overview
            if tier_context == "Tier 1" and avg_t1 is not None:
                st.markdown("**Tier 1 Matches (Top 7 vs Top 7):**")
                total_goals_t1 = (avg_t1.get("Goals Scored", 0) + avg_t1.get("Goals Conceded", 0))
                total_goals_all = (avg_all.get("Goals Scored", 0) + avg_all.get("Goals Conceded", 0))
                st.markdown(
                    f"- Average total goals per match: **{total_goals_t1:.1f}**, "
                    f"compared to **{total_goals_all:.1f}** overall."
                )
                st.markdown(
                    f"- Matches feature tighter defense with **Blocks {avg_t1.get('Blocks', 0):.1f}** and "
                    f"**GK Saves {avg_t1.get('GK Save', 0):.1f}** on average."
                )
                st.markdown(
                    "- Tier 1 contests are more structured, often decided by small efficiency margins in man-up and transition play."
                )

            elif tier_context == "Tier 2" and avg_t2 is not None:
                st.markdown("**Tier 2 Matches (Mixed / Lower-Ranked):**")
                total_goals_t2 = (avg_t2.get("Goals Scored", 0) + avg_t2.get("Goals Conceded", 0))
                total_goals_all = (avg_all.get("Goals Scored", 0) + avg_all.get("Goals Conceded", 0))
                st.markdown(
                    f"- Average total goals per match: **{total_goals_t2:.1f}**, "
                    f"compared to **{total_goals_all:.1f}** overall."
                )
                st.markdown(
                    f"- Games are more open with **Shots {avg_t2.get('Shots', 0):.1f}** and **SoT {avg_t2.get('SoT', 0):.1f}**."
                )
                st.markdown(
                    "- Tier 2 matches tend to produce more scoring chances but less defensive stability, "
                    "with discipline playing a key role."
                )

            else:
                st.markdown("**All Matches Combined:**")
                total_goals_all = (avg_all.get("Goals Scored", 0) + avg_all.get("Goals Conceded", 0))
                st.markdown(
                    f"- Average total goals per match: **{total_goals_all:.1f}**."
                )
                st.markdown(
                    f"- Typical team averages: **{avg_all.get('Shots', 0):.1f} shots**, "
                    f"**{avg_all.get('SoT', 0):.1f} SoT**, "
                    f"**{avg_all.get('Exclusions Conceded', 0):.1f} exclusions conceded**, "
                    f"and **{avg_all.get('GK Save', 0):.1f} GK saves**."
                )
                st.markdown(
                    "- Across all competitions, teams performing best tend to manage exclusion counts and maintain high defensive save efficiency."
                )

            # Optional: show Tier 1 vs Tier 2 comparison lines
            if summary_lines:
                st.markdown("#### 📊 Tier 1 vs Tier 2 Comparison Highlights:")
                for line in summary_lines:
                    st.markdown(line)

            st.markdown("<hr style='margin-top: 1rem; margin-bottom: 1rem;'>", unsafe_allow_html=True)

        except Exception as e:
            st.error(f"⚠️ Error generating context-adjusted summary: {e}")

    # -------------------------
    # Tab 1: Statistically Significant Differences (Cohen's d)
    # -------------------------
    with tabs[1]:
        try:
            # Compute Cohen's d for all numeric columns
            d_values = {col: cohen_d(df_win_filtered[col].dropna(), df_loss_filtered[col].dropna())
                        for col in num_cols}
            d_df = pd.DataFrame.from_dict(d_values, orient='index', columns=['Cohen_d']) \
                               .sort_values('Cohen_d', key=np.abs)

            st.subheader("📌 Winning and Losing Differences")

            # ➕ Coach-friendly explanation
            st.markdown("""
            **ℹ️ How to read this chart?**  
            Cohen’s d measures **how strongly a stat differs** between wins and losses.
            Which numbers change the most between wins and losses.
            - **0.2 ≈ small difference**  
            - **0.5 ≈ moderate difference**  
            - **0.8+ ≈ large difference**  
            """)

            # Bar plot with seaborn / matplotlib
            fig, ax = plt.subplots(figsize=(10, 8))
            sns.barplot(x='Cohen_d', y=d_df.index, data=d_df, ax=ax)
            ax.set_xlabel("Cohen's d Value")
            ax.set_ylabel("Statistic")
            st.pyplot(fig)

            # Optional: build dict of Cohen's d for downstream usage
            try:
                common_stats = list(set(df_win_filtered.columns) & set(df_loss_filtered.columns))
                valid_numeric_cols = df_win_filtered[common_stats].select_dtypes(include=[np.number]).columns
                ignore_cols = ['Days Since Latest', 'Time Weight']
                common_stats = [col for col in valid_numeric_cols if col not in ignore_cols]

                cohen_d_values = {}
                for stat in common_stats:
                    x = df_win_filtered[stat].dropna()
                    y = df_loss_filtered[stat].dropna()
                    if len(x) > 1 and len(y) > 1:
                        cohen_d_values[stat] = cohen_d(x, y)

                cohen_d_df = pd.DataFrame.from_dict(cohen_d_values, orient='index', columns=['Cohen_d'])
                # Store in session for reuse (e.g., Performance Index)
                st.session_state["cohen_d_df"] = cohen_d_df
            except Exception as e:
                st.error(f"Error computing Cohen's d table for reuse: {e}")

        except Exception as e:
            st.error(f"❌ Error rendering Cohen's d tab: {e}")

    # -------------------------
    # Tab 2: Performance Index
    # -------------------------
    with tabs[2]:
        st.subheader("📌 Performance Index")

        st.markdown("""
        **ℹ️ What contributes to the Performance Index?**

        The Performance Index is a weighted combination of:
        - **Win Rate** (50% weight)
        - **Average Goal Margin**: Goals Scored minus Goals Conceded (20% weight)
        - **Key Match Statistics**: Weighted by Cohen’s d values, showing which stats most distinguish wins from losses (30% weight)
        - **Tier Weighting**: Tier 1 matches (Top 7 vs Top 7) hold stronger weighting than Tier 2 matches
        """)

        with st.expander("🔝 Team Rankings by Performance Index", expanded=True):
            try:
                # Merge winning and losing data
                df_win_f = df_win_filtered.copy()
                df_loss_f = df_loss_filtered.copy()

                df_win_f['Result'] = 'Win'
                df_loss_f['Result'] = 'Lose'
                df_combined = pd.concat([df_win_f, df_loss_f], ignore_index=True)

                # --- Time Decay Weighting ---
                df_combined['Date'] = pd.to_datetime(df_combined.get('Date', pd.NaT), errors='coerce')
                df_combined = df_combined.dropna(subset=['Date'])
                latest_date = df_combined['Date'].max()
                df_combined['Days Since Latest'] = (latest_date - df_combined['Date']).dt.days
                decay_rate = 0.005  # tune if needed
                df_combined['Time Weight'] = np.exp(-decay_rate * df_combined['Days Since Latest'])

                # --- 🔥 Tier Weighting ---
                if 'Tier' in df_combined.columns:
                    tier_weights = df_combined['Tier'].map({'Tier 1': 2.0, 'Tier 2': 1.0}).fillna(1.0)
                    df_combined['Time Weight'] *= tier_weights

                # ✅ Step 1: Cohen's d (reuse if available)
                if "cohen_d_df" in st.session_state and not st.session_state["cohen_d_df"].empty:
                    cohen_d_df = st.session_state["cohen_d_df"].copy()
                else:
                    numeric_stats = df_combined.select_dtypes(include=[np.number]).columns.tolist()
                    ignore_cols = ['Days Since Latest', 'Time Weight']
                    numeric_stats = [c for c in numeric_stats if c not in ignore_cols]
                    cohen_d_values = {}
                    for stat in numeric_stats:
                        x = df_win_filtered.get(stat, pd.Series(dtype=float)).dropna()
                        y = df_loss_filtered.get(stat, pd.Series(dtype=float)).dropna()
                        if len(x) > 1 and len(y) > 1:
                            pooled_std = np.sqrt(
                                ((len(x) - 1) * x.std(ddof=1) ** 2 +
                                 (len(y) - 1) * y.std(ddof=1) ** 2) /
                                (len(x) + len(y) - 2)
                            )
                            d = (x.mean() - y.mean()) / pooled_std if pooled_std != 0 else 0
                            cohen_d_values[stat] = abs(d)
                    cohen_d_df = pd.DataFrame.from_dict(cohen_d_values, orient='index', columns=['Cohen_d'])

                # ✅ Step 2: Team-level stats with time weights
                team_stats = []
                all_teams = pd.unique(df_combined[['Winning Team', 'Losing Team']].values.ravel('K'))
                all_teams = [t for t in all_teams if pd.notna(t)]

                relevant_stats = list(cohen_d_df.index) if not cohen_d_df.empty else []

                for team in all_teams:
                    team_games = df_combined[
                        (df_combined['Winning Team'] == team) | (df_combined['Losing Team'] == team)
                        ].copy()
                    if team_games.empty:
                        continue

                    # Win Rate
                    wins = team_games[team_games['Winning Team'] == team]
                    losses = team_games[team_games['Losing Team'] == team]
                    total_games = len(wins) + len(losses)
                    win_rate = len(wins) / total_games if total_games else 0

                    # Goal Margin (use rows where the team appears; both wins and losses rows have goals)
                    margin_list = []
                    for _, row in team_games.iterrows():
                        gs = row.get('Goals Scored')
                        gc = row.get('Goals Conceded')
                        if pd.notna(gs) and pd.notna(gc):
                            margin_list.append(gs - gc)
                    avg_margin = np.mean(margin_list) if margin_list else 0

                    # Weighted average of relevant stats
                    if relevant_stats:
                        team_stat_df = team_games.copy()
                        team_stat_df['Is Team'] = (
                                (team_stat_df['Winning Team'] == team) | (team_stat_df['Losing Team'] == team)
                        )
                        team_stat_df = team_stat_df[team_stat_df['Is Team']]
                        avail_stats = [s for s in relevant_stats if s in team_stat_df.columns]
                        if avail_stats:
                            team_stat_df = team_stat_df[avail_stats + ['Time Weight']].dropna(subset=avail_stats)
                            if not team_stat_df.empty:
                                weighted_stats = (
                                    team_stat_df[avail_stats].multiply(team_stat_df['Time Weight'], axis=0)
                                ).sum()
                                total_weight = team_stat_df['Time Weight'].sum()
                                weighted_avg_stats = (
                                    weighted_stats / total_weight if total_weight > 0
                                    else pd.Series(0, index=avail_stats)
                                )
                            else:
                                weighted_avg_stats = pd.Series(0, index=avail_stats)
                        else:
                            weighted_avg_stats = pd.Series(dtype=float)
                    else:
                        weighted_avg_stats = pd.Series(dtype=float)

                    team_data = {'Team': team, 'Win Rate': win_rate, 'Goal Margin': avg_margin}
                    if not weighted_avg_stats.empty:
                        team_data.update(weighted_avg_stats.to_dict())
                    team_stats.append(team_data)

                df_perf = pd.DataFrame(team_stats).set_index('Team').fillna(0)

                # ✅ Step 3: Stat Score with “higher is better” alignment
                lower_is_better_stats = {
                    "Goals Conceded",
                    "Penalties Conceded",
                    "Missed Shots",
                    "Shots Blocked",
                    "SoT Conceded",
                    "Exclusions Conceded",
                    "Pass to CF Conceded",
                    "Pass to CF - Outcome"
                }

                if not cohen_d_df.empty:
                    rel_stats_in_perf = [s for s in cohen_d_df.index if s in df_perf.columns]
                    adjusted_stats = df_perf[rel_stats_in_perf].copy()
                    for stat in rel_stats_in_perf:
                        if stat in lower_is_better_stats:
                            max_val = df_perf[stat].max()
                            adjusted_stats[stat] = max_val - df_perf[stat]

                    # align weights
                    weights = cohen_d_df.loc[rel_stats_in_perf, 'Cohen_d']
                    df_perf['Stat Score'] = adjusted_stats.mul(weights, axis=1).sum(axis=1)
                else:
                    df_perf['Stat Score'] = 0.0

                # ✅ Step 4: Final Performance Index
                gm_div = df_perf['Goal Margin'].abs().max()
                ss_div = df_perf['Stat Score'].max()

                df_perf['Performance Index'] = (
                        df_perf['Win Rate'] * 0.5 +
                        ((df_perf['Goal Margin'] / gm_div) if gm_div and gm_div != 0 else 0) * 0.2 +
                        ((df_perf['Stat Score'] / ss_div) if ss_div and ss_div != 0 else 0) * 0.3
                )

                # ✅ Step 5: Plotly Chart
                sorted_perf = df_perf.sort_values("Performance Index", ascending=False).reset_index()

                fig = px.bar(
                    sorted_perf,
                    x="Performance Index",
                    y="Team",
                    orientation='h',
                    text="Performance Index",
                    title="🔝 Team Performance Index Ranking (Time & Tier Weighted)",
                    color="Performance Index",
                    color_continuous_scale="Blues",
                    height=600
                )
                fig.update_layout(
                    yaxis=dict(autorange="reversed"),
                    xaxis_title="Performance Index",
                    yaxis_title="Team",
                    margin=dict(l=100, r=20, t=50, b=40),
                )
                fig.update_traces(texttemplate='%{text:.2f}', textposition='outside')
                st.plotly_chart(fig, use_container_width=True)

            except Exception as e:
                st.error(f"Error calculating performance index: {e}")

    # -------------------------
    # Tab 3: Comparison Table
    # -------------------------
    with tabs[3]:
        st.subheader("📌 Team Statistical Comparison Table")

        with st.expander("🆚 Team vs Team Stat Comparison Table", expanded=True):
            try:
                # Merge win/loss into one DF with a unified 'Team' column
                df_combined_all = pd.concat([
                    df_win_filtered.rename(columns={'Winning Team': 'Team'}),
                    df_loss_filtered.rename(columns={'Losing Team': 'Team'})
                ], ignore_index=True)

                numeric_cols = df_combined_all.select_dtypes(include=[np.number]).columns

                # Average stats per team
                team_avg_stats = df_combined_all.groupby('Team')[numeric_cols].mean().dropna(how='all')
                teams = sorted(team_avg_stats.index.unique())

                team1 = st.selectbox("Select Team 1", teams, key="team1_compare_table")
                team2 = st.selectbox("Select Team 2", teams, key="team2_compare_table")

                if team1 not in team_avg_stats.index or team2 not in team_avg_stats.index:
                    st.warning("One or both selected teams do not have enough data to display a comparison.")
                    st.stop()

                # Rankings (mixed direction)
                lower_is_better_stats = {
                    "Goals Conceded",
                    "Penalties Conceded",
                    "Missed Shots",
                    "Shots Blocked",
                    "SoT Conceded",
                    "Exclusions Conceded",
                    "Pass to CF Conceded",
                    "Pass to CF - Outcome"
                }

                rankings = pd.DataFrame(index=team_avg_stats.index)
                for stat in team_avg_stats.columns:
                    if stat in lower_is_better_stats:
                        rankings[stat] = team_avg_stats[stat].rank(ascending=True)
                    else:
                        rankings[stat] = team_avg_stats[stat].rank(ascending=False)

                # Side-by-side table
                df_compare = pd.DataFrame({
                    f"{team1} Avg": team_avg_stats.loc[team1],
                    f"{team1} Rank": rankings.loc[team1],
                    "Statistic": team_avg_stats.columns,
                    f"{team2} Avg": team_avg_stats.loc[team2],
                    f"{team2} Rank": rankings.loc[team2],
                })

                df_compare = df_compare[
                    [f"{team1} Avg", f"{team1} Rank", "Statistic", f"{team2} Avg", f"{team2} Rank"]
                ].set_index("Statistic")

                # Ordinal helper
                def ordinal(n):
                    n = int(n)
                    return f"{n}{'tsnrhtdd'[(n // 10 % 10 != 1) * (n % 10 < 4) * n % 10::4]}"

                df_compare[f"{team1} Rank"] = df_compare[f"{team1} Rank"].apply(
                    lambda x: ordinal(x) if pd.notnull(x) else "")
                df_compare[f"{team2} Rank"] = df_compare[f"{team2} Rank"].apply(
                    lambda x: ordinal(x) if pd.notnull(x) else "")

                # ✅ Dynamic colouring of top 3 and bottom 3 ranks
                def highlight_top_bottom(val):
                    if isinstance(val, str) and val[:-2].isdigit():
                        rank = int(val[:-2])  # remove 'st', 'nd', 'rd', 'th'
                        total_teams = len(team_avg_stats)  # number of teams in dataset
                        if rank <= 3:
                            return 'background-color: rgba(102, 194, 165, 0.8); font-weight: bold'  # soft green
                        elif rank > total_teams - 3:
                            return 'background-color: rgba(252, 141, 98, 0.8); font-weight: bold'  # soft red
                    return ''

                styled_df = df_compare.style.applymap(
                    highlight_top_bottom, subset=[f"{team1} Rank", f"{team2} Rank"]
                ).format(precision=2)

                st.dataframe(styled_df, use_container_width=True)

            except Exception as e:
                st.error(f"❌ Error generating comparison table: {e}")

    # -------------------------
    # Tab 4: Team Breakdown
    # -------------------------
    with tabs[4]:
        try:
            team_col_win = "Winning Team"
            team_col_loss = "Losing Team"

            team_names = sorted(set(df_win_filtered[team_col_win].dropna().unique()) |
                                set(df_loss_filtered[team_col_loss].dropna().unique()))
            if not team_names:
                st.warning("No teams found after filtering.")
                st.stop()

            selected_team = st.selectbox("🔍 Select a Team for Breakdown", team_names)

            team_win_data = df_win_filtered[df_win_filtered[team_col_win] == selected_team]
            team_loss_data = df_loss_filtered[df_loss_filtered[team_col_loss] == selected_team]

            # --- Snapshot Stats ---
            total_matches = len(team_win_data) + len(team_loss_data)
            total_wins = len(team_win_data)
            total_losses = len(team_loss_data)
            win_percentage = (total_wins / total_matches * 100) if total_matches > 0 else 0

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Matches Played", total_matches)
            col2.metric("Wins", total_wins)
            col3.metric("Losses", total_losses)
            col4.metric("Win %", f"{win_percentage:.1f}%")

            if not team_win_data.empty and not team_loss_data.empty:
                # --- Average Stats Wins vs Losses
                avg_team_win = team_win_data[num_cols].mean()
                avg_team_loss = team_loss_data[num_cols].mean()

                team_df = pd.DataFrame({'Wins': avg_team_win, 'Losses': avg_team_loss}).sort_index()

                st.subheader(f"📌 {selected_team} - Average Stats in Wins vs Losses")

                team_df_reset = team_df.reset_index().rename(columns={'index': 'Statistic'})
                team_df_melted = team_df_reset.melt(
                    id_vars='Statistic', var_name='Result', value_name='Average Value'
                )

                fig = px.bar(
                    team_df_melted,
                    y='Statistic',
                    x='Average Value',
                    color='Result',
                    barmode='group',
                    orientation='h',
                    hover_name='Statistic',
                    hover_data={'Average Value': ':.2f', 'Result': True},
                    color_discrete_map={'Wins': '#66c2a5', 'Losses': '#fc8d62'},
                    height=800
                )
                fig.update_layout(xaxis_title="Average Value", yaxis_title="Statistic")
                st.plotly_chart(fig, use_container_width=True)

                # --- Boxplots of All Stats
                st.subheader(f"📌 {selected_team} - Boxplots of All Stats in Wins vs Losses")
                combined = pd.concat([
                    team_win_data[num_cols].assign(Result='Win'),
                    team_loss_data[num_cols].assign(Result='Loss')
                ])
                melted = pd.melt(combined, id_vars='Result', var_name='Statistic', value_name='Value')

                fig2 = px.box(
                    melted,
                    x="Statistic",
                    y="Value",
                    color="Result",
                    points="all",
                    title=f"{selected_team} - Boxplots of All Stats in Wins vs Losses",
                    color_discrete_map={"Win": "green", "Loss": "red"},
                    height=800
                )
                fig2.update_layout(
                    xaxis_title="Statistic",
                    yaxis_title="Value",
                    boxmode="group",
                    margin=dict(l=40, r=40, t=60, b=60)
                )
                fig2.update_traces(marker=dict(size=5, opacity=0.6))
                st.plotly_chart(fig2, use_container_width=True)
            else:
                st.warning(f"No data found for {selected_team} in one or both result categories.")

            # --- AI-style summary (exclude obvious outcome stats)
            st.subheader(f"🧠 Insights: What matters when {selected_team} win or lose")

            common_stats = list(set(team_win_data.columns) & set(team_loss_data.columns))
            valid_cols = team_win_data[common_stats].select_dtypes(include=[np.number]).columns
            ignore_cols = ['Days Since Latest', 'Time Weight', 'Goals Scored', 'Goals Conceded']
            stat_cols = [col for col in valid_cols if col not in ignore_cols]

            team_cohen_d = {}
            for stat in stat_cols:
                win_vals = team_win_data[stat].dropna()
                loss_vals = team_loss_data[stat].dropna()
                if len(win_vals) > 1 and len(loss_vals) > 1:
                    d = cohen_d(win_vals, loss_vals)
                    team_cohen_d[stat] = d

            top_stats = sorted(team_cohen_d.items(), key=lambda x: abs(x[1]), reverse=True)[:5]

            st.markdown("### 📋 Summary:")
            for stat, d in top_stats:
                avg_win = team_win_data[stat].mean()
                avg_loss = team_loss_data[stat].mean()
                better_when = "higher" if avg_win > avg_loss else "lower"
                outcome = "increased" if better_when == "higher" else "reduced"
                st.markdown(f"- When {selected_team} **win**, they tend to have **{outcome}** `{stat}` compared to when they lose.")

        except Exception as e:
            st.error(f"❌ Error in Team Breakdown tab: {e}")

    # -------------------------
    # Tab 5: Predictive Insights
    # -------------------------
    with tabs[5]:
        st.subheader("🔮 Predictive Insights")

        st.markdown("""
        **ℹ️ What this tab shows:**  
        - Correlation with Winning -> Which match statistics are most strongly linked to winning and losing.  
        - Logistic Regression -> Identifies which statistics most impact win probability.  
        - "What-If Simulation -> Sliders so you can test how changes in certain stats affect win probability.  
        """)

        try:
            # ✅ Prepare combined dataset
            df_win_f = df_win_filtered.copy()
            df_loss_f = df_loss_filtered.copy()
            df_win_f["Result"] = "Win"
            df_loss_f["Result"] = "Lose"
            df_combined = pd.concat([df_win_f, df_loss_f], ignore_index=True)

            # Add binary outcome for regression
            df_combined["WinBinary"] = (df_combined["Result"] == "Win").astype(int)

            # --- 1) Correlation with Winning (Bar Chart) ---
            st.markdown("### 📊 Correlation with Winning")

            correlations = {}
            for stat in num_cols:
                if stat in df_combined.columns:
                    try:
                        correlations[stat] = df_combined[stat].corr(df_combined["WinBinary"])
                    except Exception:
                        continue

            if correlations:
                corr_df = (
                    pd.DataFrame.from_dict(correlations, orient="index", columns=["Correlation"])
                    .dropna()
                    .sort_values("Correlation", ascending=True)
                )

                # ✅ Only show top 5 negative and top 5 positive correlations
                top_neg = corr_df.head(5)
                top_pos = corr_df.tail(5)
                corr_subset = pd.concat([top_neg, top_pos])

                fig = px.bar(
                    corr_subset,
                    x="Correlation",
                    y=corr_subset.index,
                    orientation="h",
                    title="Top Statistics Correlated with Winning",
                    color="Correlation",
                    color_continuous_scale="RdBu",
                    range_color=[-1, 1]
                )

                fig.update_layout(
                    height=500,
                    xaxis_title="Correlation with Winning",
                    yaxis_title="Statistic",
                    yaxis=dict(tickfont=dict(size=12))
                )

                st.plotly_chart(fig, use_container_width=True)

                st.markdown("""
                **ℹ️ How to read this chart:**  
                - Bars to the **right** (positive values) → Stat increases chances of winning.  
                - Bars to the **left** (negative values) → Stat decreases chances of winning.  
                - The further from 0, the stronger the relationship.  
                - Only the **Top 5 Positive** and **Top 5 Negative** correlations are shown for clarity.
                """)

            else:
                st.warning("⚠️ No valid numeric stats available for correlation analysis.")

            # --- 2) KPI Correlation Insights ---
            st.markdown("### 💡 KPI Correlation Insights")

            try:
                # Create a copy with only numeric columns and win indicator
                corr_df = df_combined[num_cols.tolist() + ["WinBinary"]].copy()
                insights = []

                for stat in num_cols:
                    if stat not in corr_df.columns:
                        continue

                    # Calculate correlation
                    corr_val = corr_df[stat].corr(corr_df["WinBinary"])
                    median_val = corr_df[stat].median()
                    mean_val = corr_df[stat].mean()

                    # Split data into above/below median groups
                    high_group = corr_df[corr_df[stat] >= median_val]
                    low_group = corr_df[corr_df[stat] < median_val]

                    # Compute win rates
                    high_win_rate = high_group["WinBinary"].mean() * 100
                    low_win_rate = low_group["WinBinary"].mean() * 100
                    diff = abs(high_win_rate - low_win_rate)

                    # Only include stats with enough variation and meaningful difference
                    if pd.notnull(corr_val) and abs(corr_val) > 0.2 and diff > 5:
                        if corr_val > 0:
                            insights.append(
                                f"📈 Teams with **above-average {stat} (avg = {mean_val:.1f})** win **{high_win_rate:.1f}%** of games, "
                                f"compared to **{low_win_rate:.1f}%** for those below average — "
                                f"a **{diff:.1f}%** difference."
                            )
                        else:
                            insights.append(
                                f"📉 Teams with **above-average {stat} (avg = {mean_val:.1f})** win only **{high_win_rate:.1f}%** of games, "
                                f"compared to **{low_win_rate:.1f}%** for those below average — "
                                f"a **{diff:.1f}%** difference."
                            )

                if insights:
                    st.markdown("#### 📊 Key Relationships Between Stats and Winning:")
                    for line in insights[:5]:  # Show top 5 most relevant
                        st.markdown(line)
                else:
                    st.info("No strong or consistent KPI correlations found in this dataset.")

            except Exception as e:
                st.error(f"⚠️ Error generating KPI correlation insights: {e}")

            # --- 3) Logistic Regression ---
            st.markdown("### 📈 Logistic Regression")
            from sklearn.linear_model import LogisticRegression
            from sklearn.preprocessing import StandardScaler

            if not num_cols.empty:
                X = df_combined[num_cols].fillna(0)
                y = df_combined["WinBinary"]

                # Scale numeric features
                scaler = StandardScaler()
                X_scaled = scaler.fit_transform(X)

                model = LogisticRegression(max_iter=500)
                model.fit(X_scaled, y)

                coefs = pd.DataFrame({
                    "Statistic": num_cols,
                    "Coefficient": model.coef_[0]
                }).sort_values("Coefficient", key=abs, ascending=False)

                st.write("Top influencing stats:")

                st.markdown("""
                **ℹ️ How to read this table:**  
                - This shows which stats have the greatest impact on winning whilst taking other statistics into account.  
                - Which statistics matter the most in predicting a win?  
                """)

                st.dataframe(
                    coefs.head(10),
                    use_container_width=False,
                )

                # --- 3) What-If Simulation ---
                st.markdown("### 🔮 What-if Simulation")
                st.markdown("""
                            **ℹ️ How to use this:**  
                            Select a team and adjust up to six match statistics to see how changes could affect their chance of winning.
                            """)

                # Select team for baseline
                teams = sorted(set(df_win_f['Winning Team'].dropna()) | set(df_loss_f['Losing Team'].dropna()))
                selected_team = st.selectbox("Select Team", teams)

                # Filter data for that team
                team_rows = df_combined[
                    (df_combined["Winning Team"] == selected_team)
                    | (df_combined["Losing Team"] == selected_team)
                    ]

                # Handle missing team data
                if team_rows.empty:
                    st.warning(f"No data found for {selected_team}. Using overall averages.")
                    team_avg = X.mean(axis=0).values
                else:
                    team_avg = team_rows[num_cols].mean().fillna(0).values

                # Baseline probability
                baseline_prob = model.predict_proba([scaler.transform([team_avg])[0]])[0, 1]
                st.write(f"Baseline predicted win probability for **{selected_team}**: **{baseline_prob:.1%}**")

                # ✅ Choose number of stats to adjust (dropdown instead of slider)
                num_adjust = st.selectbox(
                    "How many stats would you like to adjust?",
                    options=[3, 4, 5, 6],
                    index=2
                )

                # Create interactive stat selectors and sliders
                sim_inputs = {}
                for i in range(num_adjust):
                    st.markdown(f"**Adjustment {i + 1}:**")
                    col1, col2 = st.columns([1.5, 3])

                    with col1:
                        stat = st.selectbox(
                            f"Select Stat #{i + 1}",
                            options=num_cols,
                            key=f"stat_{i}"
                        )

                    with col2:
                        if stat in num_cols:
                            current_val = float(team_rows[stat].mean()) if not team_rows.empty else float(
                                df_combined[stat].mean())
                            sim_inputs[stat] = st.slider(
                                f"Adjust {stat}",
                                min_value=float(df_combined[stat].min()),
                                max_value=float(df_combined[stat].max()),
                                value=current_val,
                                step=1.0,
                                key=f"slider_{i}"
                            )

                # Apply simulation
                scenario = team_avg.copy()
                for stat, new_val in sim_inputs.items():
                    idx = list(num_cols).index(stat)
                    scenario[idx] = new_val

                # Predicted probability with adjustments
                win_prob_scenario = model.predict_proba([scaler.transform([scenario])[0]])[0, 1]

                st.success(
                    f"Predicted win probability for **{selected_team}** after adjustments: **{win_prob_scenario:.1%}**"
                )

                # --- 🧠 AI Insights Summary (robust, model-driven) ---
                st.subheader("🧠 AI Insights Summary")

                try:
                    # Build coefficient frame (fallback to coef magnitudes if available)
                    coef_df = pd.DataFrame({
                        "Statistic": list(num_cols),
                        "Coefficient": model.coef_[0] if hasattr(model, "coef_") else np.zeros(len(num_cols))
                    })

                    # Get team mean values and stat list
                    stat_list = list(num_cols)
                    team_mean_series = pd.Series(team_avg, index=stat_list)

                    # Select top 5 most influential stats by coefficient magnitude
                    top_influencers = (
                        coef_df.reindex(coef_df["Coefficient"].abs().sort_values(ascending=False).index)
                        .head(5)
                        .reset_index(drop=True)
                    )

                    st.markdown("### 📋 Key Factors Affecting Win Probability:")
                    st.markdown("""
                                    These insights test how small increases or decreases in key stats would actually change
                                    the team's win probability using the logistic regression model.
                                    """)

                    for _, row in top_influencers.iterrows():
                        stat = row["Statistic"]

                        # Determine a realistic change amount (delta)
                        if stat in df_combined.columns:
                            std_val = float(df_combined[stat].std(skipna=True))
                        else:
                            std_val = 0.0
                        delta = max(1.0, round(std_val, 1)) if std_val > 0 else 1.0

                        # Build two scenarios: +delta and -delta
                        scenario_inc = team_avg.copy()
                        scenario_dec = team_avg.copy()
                        idx = stat_list.index(stat)
                        scenario_inc[idx] += delta
                        scenario_dec[idx] -= delta

                        # Predict probabilities for each scenario
                        try:
                            prob_inc = model.predict_proba([scaler.transform([scenario_inc])[0]])[0, 1]
                            prob_dec = model.predict_proba([scaler.transform([scenario_dec])[0]])[0, 1]
                        except Exception:
                            continue

                        # Compare both directions against baseline
                        if prob_inc > baseline_prob or prob_dec > baseline_prob:
                            if prob_inc >= prob_dec:
                                best_action = "increase"
                                new_prob = prob_inc
                                change = prob_inc - baseline_prob
                                change_val = delta
                            else:
                                best_action = "decrease"
                                new_prob = prob_dec
                                change = prob_dec - baseline_prob
                                change_val = delta

                            trend = "↑" if change > 0 else "↓"
                            st.markdown(
                                f"- If **{selected_team}** {best_action}s their **{stat}** by **{change_val}**, "
                                f"predicted win probability changes from **{baseline_prob:.1%} → {new_prob:.1%}** "
                                f"({trend} {abs(change) * 100:.1f}%)."
                            )
                        else:
                            st.markdown(
                                f"- **{stat}** — Adjusting this stat by ±{delta} has minimal short-term effect on win probability."
                            )

                    # Identify single biggest lever
                    best_candidate = None
                    best_increase = 0.0
                    for _, row in top_influencers.iterrows():
                        stat = row["Statistic"]
                        idx = stat_list.index(stat)
                        delta = 1.0
                        scenario_inc = team_avg.copy()
                        scenario_dec = team_avg.copy()
                        scenario_inc[idx] += delta
                        scenario_dec[idx] -= delta

                        try:
                            prob_inc = model.predict_proba([scaler.transform([scenario_inc])[0]])[0, 1]
                            prob_dec = model.predict_proba([scaler.transform([scenario_dec])[0]])[0, 1]
                        except Exception:
                            continue

                        best_local = max(prob_inc, prob_dec)
                        increase = best_local - baseline_prob
                        if increase > best_increase:
                            best_increase = increase
                            best_candidate = (stat, increase, baseline_prob, best_local)

                    if best_candidate is not None and best_increase > 0:
                        stat, increase, base, newp = best_candidate
                        st.info(
                            f"📌 Biggest lever: changing **{stat}** slightly could raise "
                            f"**{selected_team}**’s win probability by **{increase * 100:.1f}%** "
                            f"(from {base:.1%} → {newp:.1%})."
                        )

                except Exception as e:
                    st.error(f"⚠️ Could not generate AI insights: {e}")

            else:
                st.warning("⚠️ No numeric columns found for regression or simulation.")

        except Exception as e:
            st.error(f"❌ Error generating predictive insights: {e}")

        # -------------------------
        # Tab 6: Referee Analysis
        # -------------------------
        with tabs[6]:
            st.subheader("⚖️ Referee Analysis")

            st.markdown("""
            **ℹ️ What this shows:**  
            This section analyses match trends based on referees, highlighting how the flow of a game  
            (goals, penalties, exclusions, and centre-forward outcomes) varies depending on who officiated.
            """)

            try:
                # ✅ Use ONLY Winning Teams sheet to avoid duplicate matches
                df_ref = df_win_filtered.copy()

                # Ensure referee columns exist
                if 'Referee 1' not in df_ref.columns or 'Referee 2' not in df_ref.columns:
                    st.warning("⚠️ Referee columns not found in the dataset.")
                else:
                    # Create a combined referee list
                    referees = sorted(set(df_ref['Referee 1'].dropna()) | set(df_ref['Referee 2'].dropna()))
                    selected_ref = st.selectbox("Select Referee", referees)

                    # Filter matches where this referee appears
                    ref_matches = df_ref[
                        (df_ref['Referee 1'] == selected_ref) | (df_ref['Referee 2'] == selected_ref)
                        ].copy()

                    if ref_matches.empty:
                        st.warning(f"No matches found for {selected_ref}.")
                    else:
                        # ✅ Compute total stats per match
                        ref_matches["Total Goals"] = ref_matches["Goals Scored"].fillna(0) + ref_matches[
                            "Goals Conceded"].fillna(0)
                        ref_matches["Total Penalties"] = ref_matches["Penalties Awarded"].fillna(0) + ref_matches[
                            "Penalties Conceded"].fillna(0)
                        ref_matches["Total Exclusions"] = ref_matches["Exclusions Won"].fillna(0) + ref_matches[
                            "Exclusions Conceded"].fillna(0)

                        # Dataset-wide averages for context
                        global_avg_goals = (
                                    df_ref["Goals Scored"].fillna(0) + df_ref["Goals Conceded"].fillna(0)).mean()
                        global_avg_penalties = (
                                    df_ref["Penalties Awarded"].fillna(0) + df_ref["Penalties Conceded"].fillna(
                                0)).mean()
                        global_avg_exclusions = (
                                    df_ref["Exclusions Won"].fillna(0) + df_ref["Exclusions Conceded"].fillna(0)).mean()

                        # Referee-specific averages
                        avg_goals = ref_matches["Total Goals"].mean()
                        avg_penalties = ref_matches["Total Penalties"].mean()
                        avg_exclusions = ref_matches["Total Exclusions"].mean()

                        st.markdown(f"**Referee:** {selected_ref} — Matches officiated: **{len(ref_matches)}**")

                        col1, col2, col3 = st.columns(3)
                        col1.metric("⚽ Avg Total Goals per Match", f"{avg_goals:.1f}",
                                    f"Global Avg: {global_avg_goals:.1f}")
                        col2.metric("🚩 Avg Total Penalties per Match", f"{avg_penalties:.1f}",
                                    f"Global Avg: {global_avg_penalties:.1f}")
                        col3.metric("❌ Avg Total Exclusions per Match", f"{avg_exclusions:.1f}",
                                    f"Global Avg: {global_avg_exclusions:.1f}")

                        # 🎯 Centre-Forward Pass Outcomes
                        st.markdown("### 🎯 Centre-Forward Pass Outcomes")

                        # Define expected CF columns
                        cf_cols = [
                            "Pass to CF + Outcome",
                            "Pass to CF Neutral Outcome",
                            "Pass to CF - Outcome"
                        ]
                        cf_data = {}
                        global_cf_data = {}

                        for col in cf_cols:
                            if col in ref_matches.columns:
                                cf_data[col] = ref_matches[col].fillna(0).mean()
                                global_cf_data[col] = df_ref[col].fillna(0).mean()

                        if cf_data:
                            # Order so +, Neutral, - appears in that order
                            cf_order = ["Pass to CF + Outcome", "Pass to CF Neutral Outcome", "Pass to CF - Outcome"]
                            cf_df = pd.DataFrame({
                                "Outcome": [c for c in cf_order if c in cf_data],
                                "Avg per match": [cf_data[c] for c in cf_order if c in cf_data],
                                "Global Avg": [global_cf_data[c] for c in cf_order if c in global_cf_data]
                            })
                            cf_df["Outcome"] = cf_df["Outcome"].replace({
                                "Pass to CF + Outcome": "Pass to CF +",
                                "Pass to CF Neutral Outcome": "Pass to CF Neutral",
                                "Pass to CF - Outcome": "Pass to CF -"
                            })

                            # Custom color map for outcomes
                            color_map = {
                                "Pass to CF +": "#2ecc71",  # green
                                "Pass to CF Neutral": "#f39c12",  # orange
                                "Pass to CF -": "#e74c3c"  # red
                            }

                            col_chart, col_avg = st.columns([3, 1])

                            with col_chart:
                                fig = px.bar(
                                    cf_df,
                                    x="Avg per match",
                                    y="Outcome",
                                    orientation="h",
                                    title=f"Centre-Forward Pass Outcomes for {selected_ref}",
                                    text="Avg per match",
                                    color="Outcome",
                                    color_discrete_map=color_map
                                )
                                fig.update_traces(texttemplate='%{text:.2f}', textposition='outside')
                                fig.update_layout(
                                    yaxis_title="Outcome",
                                    xaxis_title="Average per Match",
                                    height=320,
                                    margin=dict(l=120, r=20, t=40, b=40),
                                    showlegend=False
                                )
                                st.plotly_chart(fig, use_container_width=True)

                            with col_avg:
                                st.markdown("#### ⚖️ Global Averages")
                                for _, row in cf_df.iterrows():
                                    st.markdown(f"**{row['Outcome']}**: {row['Global Avg']:.2f}")

                        else:
                            st.info(f"No 'Pass to CF' data available for {selected_ref}.")

                        # 🏆 Team-specific win/loss record under this referee
                        st.markdown("### 🏅 Team Results under this Referee")

                        # Only proceed if 'Winning Team' exists
                        if "Winning Team" in ref_matches.columns:
                            # Try to also pull 'Losing Team' if available
                            losing_available = "Losing Team" in df_loss_filtered.columns

                            team_results = []
                            all_teams = set(ref_matches['Winning Team'].dropna().unique())
                            if losing_available:
                                all_teams |= set(df_loss_filtered['Losing Team'].dropna().unique())

                            for team in sorted(all_teams):
                                wins = len(ref_matches[ref_matches['Winning Team'] == team])
                                losses = 0
                                if losing_available:
                                    losses = len(df_loss_filtered[
                                                     ((df_loss_filtered['Referee 1'] == selected_ref) |
                                                      (df_loss_filtered['Referee 2'] == selected_ref)) &
                                                     (df_loss_filtered['Losing Team'] == team)
                                                     ])
                                total = wins + losses
                                if total > 0:
                                    team_results.append(
                                        {"Team": team, "Wins": wins, "Losses": losses, "Matches": total})

                            if team_results:
                                team_summary = pd.DataFrame(team_results).sort_values("Team")
                                team_summary["Win %"] = (team_summary["Wins"] / team_summary["Matches"] * 100).round(1)
                                st.dataframe(
                                    team_summary[["Team", "Matches", "Wins", "Losses", "Win %"]],
                                    use_container_width=True,
                                    hide_index=True
                                )
                            else:
                                st.info("No team results available for this referee.")
                        else:
                            st.info("Winning Team column not found — cannot calculate team results.")

            except Exception as e:
                st.error(f"❌ Error generating Referee Analysis: {e}")

            # -------------------------
            # Tab 7: Opponent Profiling
            # -------------------------
            with tabs[7]:
                st.subheader("🎯 Opponent Profiling")

                st.markdown("""
                **ℹ️ What this shows:**  
                A tactical and statistical fingerprint of the selected opponent — including their key strengths, weaknesses,
                and style indicators relative to all teams.
                """)

                try:
                    # Combine filtered data (already filtered by competition / tier / timeframe)
                    df_combined_opp = pd.concat([df_win_filtered, df_loss_filtered], ignore_index=True)

                    # Build set of all teams (from both columns)
                    all_teams = sorted(
                        set(df_combined_opp.get('Winning Team', pd.Series(dtype=str)).dropna()) |
                        set(df_combined_opp.get('Losing Team', pd.Series(dtype=str)).dropna())
                    )

                    if not all_teams:
                        st.warning("No teams found in the dataset.")
                        st.stop()

                    selected_team = st.selectbox("Select Opponent to Profile", all_teams)

                    # --- Build per-team average table for numeric KPIs (one row per team) ---
                    # Numeric columns to consider
                    numeric_cols = df_combined_opp.select_dtypes(include=[np.number]).columns.tolist()
                    if not numeric_cols:
                        st.warning("No numeric columns found for profiling.")
                        st.stop()

                    # Function to compute team-average for a given team (works if team appears as winning or losing)
                    def team_mean_for(team, cols):
                        rows = df_combined_opp[
                            (df_combined_opp.get('Winning Team') == team) | (df_combined_opp.get('Losing Team') == team)
                            ]
                        if rows.empty:
                            return pd.Series({c: np.nan for c in cols})
                        return rows[cols].mean()

                    # Build DataFrame of team averages
                    team_averages = []
                    for t in all_teams:
                        s = team_mean_for(t, numeric_cols)
                        s.name = t
                        team_averages.append(s)
                    team_avg_df = pd.DataFrame(team_averages)

                    # Global average (for context)
                    global_avg = df_combined_opp[numeric_cols].mean()

                    # KPIs where lower values indicate better performance
                    lower_is_better = {
                        "Goals Conceded", "Penalties Conceded", "Missed Shots",
                        "Shots Blocked", "SoT Conceded", "Exclusions Conceded",
                        "Pass to CF Conceded", "Pass to CF - Outcome"
                    }

                    # Compute directional difference for selected team relative to global average:
                    # For 'higher is better': diff = team - global  (bigger positive = strength)
                    # For 'lower is better':  diff = global - team  (bigger positive = strength)
                    team_row = team_avg_df.loc[selected_team]
                    direction_diff = {}
                    for stat in numeric_cols:
                        tval = team_row.get(stat, np.nan)
                        gval = global_avg.get(stat, np.nan)
                        if pd.isna(tval) or pd.isna(gval):
                            direction_diff[stat] = np.nan
                        else:
                            if stat in lower_is_better:
                                direction_diff[stat] = (gval - tval)  # positive => team is better (concedes fewer)
                            else:
                                direction_diff[stat] = (tval - gval)  # positive => team is better (scores more)

                    diff_series = pd.Series(direction_diff).dropna().sort_values(ascending=False)

                    # Top 3 strengths (largest positive diffs) and top 3 weaknesses (most negative diffs)
                    top_strengths = diff_series.head(3)
                    top_weaknesses = diff_series.tail(3).sort_values(ascending=True)  # show worst first

                    # --- Helper: ordinal formatter ---
                    def ordinal(n):
                        n = int(n)
                        if 10 <= n % 100 <= 20:
                            suf = "th"
                        else:
                            suf = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
                        return f"{n}{suf}"

                    # --- Precompute ranks per stat (1 = best). For lower_is_better stats, smaller is better so rank ascending.
                    rank_df = pd.DataFrame(index=team_avg_df.index)
                    for stat in numeric_cols:
                        series = team_avg_df[stat]
                        if stat in lower_is_better:
                            # lower -> better => rank ascending (1 = lowest)
                            rank_df[stat] = series.rank(ascending=True, method="min")
                        else:
                            # higher -> better => rank descending (1 = highest)
                            rank_df[stat] = series.rank(ascending=False, method="min")

                    # Total teams for contextualizing rank
                    total_teams = len(team_avg_df)

                    # --- Display header summary ---
                    total_matches = len(df_combined_opp[
                                            (df_combined_opp.get('Winning Team') == selected_team) |
                                            (df_combined_opp.get('Losing Team') == selected_team)
                                            ])
                    wins = int((df_combined_opp.get('Winning Team') == selected_team).sum())
                    win_rate = (wins / total_matches) * 100 if total_matches > 0 else 0.0

                    st.markdown(f"## 🏐 {selected_team} Profile")
                    c1, c2 = st.columns(2)
                    c1.metric("📊 Matches Analysed", f"{total_matches}")
                    c2.metric("🏆 Win %", f"{win_rate:.1f}%")

                    import plotly.graph_objects as go

                    gauge_height = 170  # slightly smaller so titles are readable

                    # --- Strengths section ---
                    st.markdown("### 💪 Strengths")
                    if top_strengths.empty:
                        st.info("No clear strengths (not enough data).")
                    else:
                        st.caption("Metrics where the team outperforms the competition average (rank shown).")
                        for stat, diff_val in top_strengths.items():
                            tval = team_row.get(stat, np.nan)
                            gval = global_avg.get(stat, np.nan)
                            rank_val = int(rank_df.loc[selected_team, stat]) if stat in rank_df.columns else None
                            rank_text = f" ({ordinal(rank_val)}/{total_teams})" if pd.notna(rank_val) else ""
                            left, right = st.columns([1, 2])
                            with left:
                                st.markdown(f"**{stat}**")
                                st.markdown(f"- Team Avg: **{tval:.2f}**{rank_text}  \n- Global Avg: **{gval:.2f}**")
                            with right:
                                # Axis range: anchor at zero up to a sensible max around team/global
                                axis_max = max(gval * 1.5 if gval > 0 else tval * 2, tval * 1.2 if tval > 0 else 1)
                                if axis_max <= 0:
                                    axis_max = 1
                                # green gauge for strengths
                                fig = go.Figure(go.Indicator(
                                    mode="gauge+number+delta",
                                    value=round(float(tval) if pd.notna(tval) else 0, 2),
                                    delta={'reference': round(float(gval) if pd.notna(gval) else 0, 2),
                                           'increasing': {'color': "green"}, 'decreasing': {'color': "gray"}},
                                    gauge={
                                        'axis': {'range': [0, axis_max]},
                                        'bar': {'color': "seagreen"},
                                        'steps': [
                                            {'range': [0, gval], 'color': "#f0f0f0"},
                                            {'range': [gval, axis_max], 'color': "#d4efdf"}
                                        ],
                                        'threshold': {'line': {'color': "black", 'width': 2}, 'value': gval}
                                    },
                                    title={'text': f"{stat}", 'font': {'size': 12}}
                                ))
                                fig.update_layout(height=gauge_height, margin=dict(t=10, b=0, l=0, r=0))
                                st.plotly_chart(fig, use_container_width=True)

                    st.markdown("---")

                    # --- Weaknesses section ---
                    st.markdown("### ⚠️ Weaknesses")
                    if top_weaknesses.empty:
                        st.info("No clear weaknesses (not enough data).")
                    else:
                        st.caption("Metrics where the team performs worse than the competition average (rank shown).")
                        for stat, diff_val in top_weaknesses.items():
                            tval = team_row.get(stat, np.nan)
                            gval = global_avg.get(stat, np.nan)
                            rank_val = int(rank_df.loc[selected_team, stat]) if stat in rank_df.columns else None
                            rank_text = f" ({ordinal(rank_val)}/{total_teams})" if pd.notna(rank_val) else ""
                            left, right = st.columns([1, 2])
                            with left:
                                st.markdown(f"**{stat}**")
                                st.markdown(f"- Team Avg: **{tval:.2f}**{rank_text}  \n- Global Avg: **{gval:.2f}**")
                            with right:
                                # Axis range and red colour for weakness
                                axis_max = max(gval * 1.5 if gval > 0 else tval * 2, tval * 1.2 if tval > 0 else 1)
                                if axis_max <= 0:
                                    axis_max = 1
                                fig = go.Figure(go.Indicator(
                                    mode="gauge+number+delta",
                                    value=round(float(tval) if pd.notna(tval) else 0, 2),
                                    delta={'reference': round(float(gval) if pd.notna(gval) else 0, 2),
                                           'increasing': {'color': "gray"}, 'decreasing': {'color': "red"}},
                                    gauge={
                                        'axis': {'range': [0, axis_max]},
                                        'bar': {'color': "crimson"},
                                        'steps': [
                                            {'range': [0, gval], 'color': "#fdecea"},
                                            {'range': [gval, axis_max], 'color': "#f9e5e5"}
                                        ],
                                        'threshold': {'line': {'color': "black", 'width': 2}, 'value': gval}
                                    },
                                    title={'text': f"{stat}", 'font': {'size': 12}}
                                ))
                                fig.update_layout(height=gauge_height, margin=dict(t=10, b=0, l=0, r=0))
                                st.plotly_chart(fig, use_container_width=True)

                    st.markdown("---")

                    # --- Fingerprint Radar Chart (category-level aggregation) ---
                    st.markdown("### 🕸️ Team Fingerprint Summary")

                    # Define radar categories and which stats feed them (ensure stats exist)
                    radar_metrics = {
                        "Offense": ["Goals Scored", "6v6 Goals", "SoT"],
                        "Defense": ["Goals Conceded", "Shots Blocked", "GK Save"],
                        "Discipline": ["Exclusions Conceded", "Penalties Conceded"],
                        "Efficiency": ["SoT", "Shots"],
                        "CF Play": ["Pass to CF + Outcome", "Pass to CF - Outcome"]
                    }

                    radar_team = {}
                    radar_global = {}

                    for cat, stats in radar_metrics.items():
                        team_vals = []
                        global_vals = []
                        for s in stats:
                            if s in team_avg_df.columns:
                                team_vals.append(team_avg_df.loc[selected_team, s])
                                global_vals.append(global_avg[s])
                        # if no stats available for category, fallback to neutral 50
                        if len(team_vals) == 0:
                            radar_team[cat] = 50.0
                            radar_global[cat] = 50.0
                            continue

                        # For categories containing 'lower is better' stats, invert relative comparison
                        # Compute a normalized score around 50 where >50 is better than average
                        ratios = []
                        for t_val, g_val, s in zip(team_vals, global_vals, stats):
                            if pd.isna(t_val) or pd.isna(g_val) or g_val == 0:
                                ratio_score = 50.0
                            else:
                                if s in lower_is_better:
                                    # If lower is better, less than global => >50 score
                                    ratio_score = 50 + ((g_val - t_val) / g_val) * 50
                                else:
                                    ratio_score = 50 + ((t_val - g_val) / g_val) * 50
                            ratios.append(np.clip(ratio_score, 0, 100))
                        radar_team[cat] = float(np.nanmean(ratios))
                        radar_global[cat] = 50.0

                    radar_df = pd.DataFrame({
                        "Category": list(radar_team.keys()),
                        "Team": list(radar_team.values()),
                        "Global Avg": list(radar_global.values())
                    })

                    # Choose radar color dynamically
                    avg_perf = np.nanmean(list(radar_team.values()))
                    radar_color = "seagreen" if avg_perf >= 60 else ("gold" if avg_perf >= 45 else "crimson")

                    fig = go.Figure()
                    fig.add_trace(go.Scatterpolar(
                        r=radar_df["Team"],
                        theta=radar_df["Category"],
                        fill='toself',
                        name=selected_team,
                        line=dict(color=radar_color)
                    ))
                    fig.add_trace(go.Scatterpolar(
                        r=radar_df["Global Avg"],
                        theta=radar_df["Category"],
                        fill='toself',
                        name="Global Avg",
                        line=dict(color="lightgray", dash="dash")
                    ))
                    fig.update_layout(
                        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                        showlegend=True,
                        height=480,
                        title=f"Overall Performance Fingerprint: {selected_team}"
                    )
                    st.plotly_chart(fig, use_container_width=True)

                    # -------------------------
                    # REPORT EXPORT (Opponent Report) - robust DOCX with optional PDF conversion
                    # -------------------------
                    st.markdown("### 📄 Export Opponent Report")

                    import tempfile
                    import os
                    from docx import Document

                    def create_opponent_report_docx(team_name, total_matches, win_rate, strengths, weaknesses, team_avg,
                                                    global_avg):
                        doc = Document()
                        doc.add_heading(f'{team_name} — Opponent Scouting Report', level=1)

                        doc.add_heading('Summary', level=2)
                        doc.add_paragraph(f"Matches Analysed: {total_matches}")
                        doc.add_paragraph(f"Win Percentage: {win_rate:.1f}%")

                        doc.add_heading('Top Strengths', level=2)
                        for stat in strengths.index:
                            tval = team_avg.get(stat, float('nan'))
                            gval = global_avg.get(stat, float('nan'))
                            doc.add_paragraph(f"- {stat}: Team Avg {tval:.2f}, Global Avg {gval:.2f}")

                        doc.add_heading('Top Weaknesses', level=2)
                        for stat in weaknesses.index:
                            tval = team_avg.get(stat, float('nan'))
                            gval = global_avg.get(stat, float('nan'))
                            doc.add_paragraph(f"- {stat}: Team Avg {tval:.2f}, Global Avg {gval:.2f}")

                        return doc

                    try:
                        # Build strength/weakness series (reuse variables from Opponent Profile section)
                        strength_series = top_strengths if 'top_strengths' in locals() else pd.Series()
                        weakness_series = top_weaknesses if 'top_weaknesses' in locals() else pd.Series()
                        team_row = team_avg if 'team_avg' in locals() else (
                            pd.Series() if 'team_avg' not in globals() else globals()['team_avg'])
                        global_row = global_avg if 'global_avg' in locals() else (
                            pd.Series() if 'global_avg' not in globals() else globals()['global_avg'])

                        # create temporary docx
                        doc = create_opponent_report_docx(selected_team, total_matches, win_rate, strength_series,
                                                          weakness_series, team_row, global_row)
                        docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
                        doc.save(docx_path)

                        # Try convert to PDF if pypandoc available (may fail if pandoc binary missing)
                        pdf_produced = False
                        try:
                            import pypandoc
                            pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
                            # conversion may raise if pandoc binary not present
                            pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
                            pdf_produced = os.path.exists(pdf_path)
                        except Exception as conv_e:
                            pdf_produced = False

                        # Provide download buttons
                        if pdf_produced:
                            with open(pdf_path, "rb") as f:
                                st.download_button(
                                    label="📥 Download PDF Report",
                                    data=f,
                                    file_name=f"{selected_team}_Opponent_Report.pdf",
                                    mime="application/pdf"
                                )

                        # Always offer DOCX fallback
                        with open(docx_path, "rb") as f:
                            st.download_button(
                                label="📥 Download DOCX Report (recommended)",
                                data=f,
                                file_name=f"{selected_team}_Opponent_Report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                    except Exception as e:
                        st.error(f"⚠️ Could not create report file: {e}")
                    finally:
                        # cleanup temp files if they exist (wrapped in try to avoid masking errors)
                        try:
                            if 'docx_path' in locals() and os.path.exists(docx_path):
                                os.remove(docx_path)
                            if 'pdf_path' in locals() and os.path.exists(pdf_path):
                                os.remove(pdf_path)
                        except Exception:
                            pass

                except Exception as e:
                    st.error(f"❌ Error generating Opponent Profile: {e}")

if __name__ == "__main__":
    main()